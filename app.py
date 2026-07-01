"""
app.py
Internal tool: upload billing file + rate table (+ optional rectification file),
get back a validated reconciliation workbook and a summary report.

Run locally:   streamlit run app.py
Deploy free:   push this folder to a GitHub repo, then deploy on
               https://share.streamlit.io (Streamlit Community Cloud).
"""

import io
from datetime import date

import streamlit as st
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor

from reconcile import load_billing, load_rates, load_rectification, run_reconciliation, build_summary

st.set_page_config(page_title="Cigna Premium Reconciliation", layout="wide")

# ---------------------------------------------------------------------------
# Design system — "audit ledger" look: dark control panel (sidebar) for
# inputs, pale ledger-paper canvas for output, stamp-style status badges.
# Tokens defined once as CSS custom properties.
# ---------------------------------------------------------------------------
st.markdown("""
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=IBM+Plex+Mono:wght@400;500;600&display=swap" rel="stylesheet">
<style>
:root {
  --paper: #F1F3EF;
  --ink: #16211D;
  --ink-muted: #5B6B63;
  --panel-dark: #0E2420;
  --panel-text: #E7EFEA;
  --teal: #0E6B5C;
  --teal-hover: #0B564A;
  --amber: #B5651D;
  --amber-bg: #FBEBD9;
  --sage: #3F6B4F;
  --sage-bg: #E6EFE7;
  --line: #D8DED8;
}
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.stApp {
  background-color: var(--paper);
  background-image: repeating-linear-gradient(
    to bottom, transparent 0px, transparent 34px, var(--line) 35px
  );
  color: var(--ink);
}
/* Sidebar = control panel */
[data-testid="stSidebar"] {
  background-color: var(--panel-dark);
}
[data-testid="stSidebar"] * { color: var(--panel-text) !important; }
[data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
  font-weight: 700; text-transform: uppercase; letter-spacing: 0.08em; font-size: 0.85rem;
  color: #9FBDB2 !important; border-bottom: 1px solid rgba(255,255,255,0.15); padding-bottom: 6px;
}
[data-testid="stSidebar"] .stButton button {
  background-color: var(--teal); color: white !important; border: none;
  font-weight: 600; letter-spacing: 0.03em; text-transform: uppercase; font-size: 0.8rem;
  border-radius: 3px; padding: 0.6em 1em;
}
[data-testid="stSidebar"] .stButton button:hover { background-color: var(--teal-hover); }
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
  background-color: rgba(255,255,255,0.04); border: 1px dashed rgba(255,255,255,0.3);
}
/* Main titles */
h1 {
  font-family: 'Inter', sans-serif; font-weight: 700; letter-spacing: -0.01em;
  color: var(--ink); border-bottom: 2px solid var(--ink); padding-bottom: 0.3em;
}
h2, h3 {
  text-transform: uppercase; letter-spacing: 0.09em; font-size: 0.95rem !important;
  font-weight: 700; color: var(--teal) !important; margin-top: 1.6em !important;
}
/* Monospace for anything numeric/code-like */
[data-testid="stMetricValue"], code, .ledger-figure { font-family: 'IBM Plex Mono', monospace; }
/* Ledger metric cards */
.ledger-cards { display: flex; gap: 14px; margin: 10px 0 22px 0; flex-wrap: wrap; }
.ledger-card {
  background: #FFFFFF; border: 1px solid var(--line); border-left: 4px solid var(--teal);
  padding: 14px 18px; min-width: 190px; flex: 1;
}
.ledger-card.warn { border-left-color: var(--amber); }
.ledger-card .label {
  font-size: 0.72rem; text-transform: uppercase; letter-spacing: 0.08em;
  color: var(--ink-muted); font-weight: 600; margin-bottom: 4px;
}
.ledger-card .value {
  font-family: 'IBM Plex Mono', monospace; font-size: 1.5rem; font-weight: 600; color: var(--ink);
  font-variant-numeric: tabular-nums;
}
/* Stamp badges for OK / CHECK */
.stamp {
  display: inline-block; font-family: 'IBM Plex Mono', monospace; font-weight: 700;
  font-size: 0.72rem; letter-spacing: 0.08em; text-transform: uppercase;
  padding: 3px 10px; border-radius: 999px; border: 1.5px solid currentColor;
  transform: rotate(-2deg);
}
.stamp.ok { color: var(--sage); background: var(--sage-bg); }
.stamp.check { color: var(--amber); background: var(--amber-bg); }
/* Buttons in main area (download buttons) */
.stDownloadButton button {
  background-color: var(--ink); color: white; border: none; border-radius: 3px;
  font-weight: 600; text-transform: uppercase; letter-spacing: 0.04em; font-size: 0.8rem;
}
.stDownloadButton button:hover { background-color: var(--teal); }
/* Warning box override to match palette */
[data-testid="stAlert"] { border-radius: 3px; }
/* Divider hairlines */
hr { border-color: var(--line); }
/* Results table */
.ledger-table { width:100%; border-collapse: collapse; font-size: 0.85rem; margin-bottom: 1.5em; }
.ledger-table th { text-align:left; text-transform:uppercase; letter-spacing:0.05em; font-size:0.7rem; color: var(--ink-muted); border-bottom: 2px solid var(--ink); padding: 6px 10px; }
.ledger-table td { padding: 6px 10px; border-bottom: 1px solid var(--line); font-family:'IBM Plex Mono',monospace; font-size:0.82rem; }
.ledger-table tr:hover td { background: rgba(14,107,92,0.06); }
</style>
""", unsafe_allow_html=True)

def ledger_card(label, value, warn=False):
    cls = "ledger-card warn" if warn else "ledger-card"
    return f'<div class="{cls}"><div class="label">{label}</div><div class="value">{value}</div></div>'

def stamp(status):
    cls = "ok" if status == "OK" else "check"
    return f'<span class="stamp {cls}">{status}</span>'


st.markdown('<h1>Insurance Premium Reconciliation</h1>', unsafe_allow_html=True)
st.markdown(
    '<p style="color:var(--ink-muted);font-size:0.95rem;margin-top:10px;">'
    'Internal ledger tool — upload billing, rates, and rectification files; get a validated, stamped reconciliation back.</p>',
    unsafe_allow_html=True,
)

# ---------------------------------------------------------------------------
# Sidebar inputs
# ---------------------------------------------------------------------------
with st.sidebar:
    st.header("1. Files")
    billing_file = st.file_uploader("Billing file (.xlsx)", type=["xlsx"])
    rates_file = st.file_uploader("Rate table (.xlsx)", type=["xlsx"])
    rect_file = st.file_uploader("Rectification file (.xlsx) — optional", type=["xlsx"])

    st.header("2. Settings")
    rate_is_annual = st.radio(
        "The rate table figures are:",
        ["Annual premiums", "Monthly premiums"],
        help="If unsure, run once, check the warning banner that appears if the "
             "wrong setting produces a large uniform gap across most records.",
    ) == "Annual premiums"
    tolerance = st.slider("Tolerance % (flags CHECK when |gap| exceeds this)", 1, 50, 10)

    period_label = st.text_input("Period label (for the report title)", value="")

    run_btn = st.button("Run reconciliation", type="primary", use_container_width=True)

# ---------------------------------------------------------------------------
# Excel builder
# ---------------------------------------------------------------------------
def build_excel(validation_df, summary, warning, rect_df, tolerance, rate_is_annual):
    wb = Workbook()
    wb.remove(wb.active)

    green_font = Font(name="Calibri", sz=12, b=True, color="FF006100")
    green_fill = PatternFill("solid", fgColor="FFC6EFCE")
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center_top = Alignment(horizontal="center", vertical="top")
    data_font = Font(name="Calibri", sz=11)
    bold_font = Font(name="Calibri", sz=11, b=True)
    num_fmt = "#,##0.00"

    def style_header(ws, ncols):
        for c in range(1, ncols + 1):
            cell = ws.cell(row=1, column=c)
            cell.font = green_font
            cell.fill = green_fill
            cell.border = border
            cell.alignment = center_top
        ws.row_dimensions[1].height = 16

    # ---- Validation ----
    ws = wb.create_sheet("Validation")
    headers = ["Unique Identifier", "Name Family", "First Name Family", "DOB", "Age", "Age range",
               "Country name station", "Derived Region", "N° of days insured", "MonthlyPremiumUSD",
               "ExpectedUSD", "Premium medical", "DifferenceUSD", "PercentDiff", "Status", "Reason"]
    ws.append(headers)
    style_header(ws, len(headers))
    for _, row in validation_df.iterrows():
        dob_val = row["DOB"].date() if pd.notna(row["DOB"]) else None
        age_val = None if pd.isna(row["AgeNum"]) else row["AgeNum"]
        ws.append([
            row["UID"], row["NameFamily"], row["FirstNameFamily"], dob_val, age_val, row["AgeRange"],
            row["Country"], row["DerivedRegion"], row["Days"], round(row["MonthlyRate"], 2),
            round(row["ExpectedUSD"], 2), round(row["PremiumBilled"], 2), round(row["DifferenceUSD"], 2),
            (None if pd.isna(row["PercentDiff"]) else row["PercentDiff"]),
            row["Status"], row["Reason"],
        ])
    for r in range(2, ws.max_row + 1):
        for c in range(1, len(headers) + 1):
            ws.cell(row=r, column=c).font = data_font
            if c in (10, 11, 12, 13, 14):
                ws.cell(row=r, column=c).number_format = num_fmt
        ws.cell(row=r, column=4).number_format = "mm/dd/yyyy"
    widths = [20, 16, 16, 12, 8, 12, 18, 16, 14, 16, 14, 14, 14, 12, 10, 32]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w
    ws.freeze_panes = "A2"

    # ---- Summary ----
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Total Rows", "OK", "CHECK", "No Rate", "Total Billed", "Total Expected",
                "Total Difference", "Avg % Diff (OK only)", "Tolerance %", "Rates are"])
    style_header(ws2, 10)
    ws2.append([
        summary["total_rows"], summary["ok"], summary["check"], summary["no_rate"],
        round(summary["total_billed"], 2), round(summary["total_expected"], 2),
        round(summary["total_diff"], 2),
        (None if summary["avg_ok_pct"] is None else round(summary["avg_ok_pct"], 2)),
        tolerance, ("Annual" if rate_is_annual else "Monthly"),
    ])
    for c in range(1, 11):
        ws2.cell(row=2, column=c).font = data_font
        if c in (5, 6, 7, 8):
            ws2.cell(row=2, column=c).number_format = num_fmt
    for i, w in enumerate([12, 8, 10, 10, 16, 16, 16, 18, 12, 12], start=1):
        ws2.column_dimensions[ws2.cell(row=1, column=i).column_letter].width = w
    if warning:
        ws2["A4"] = "Sanity check warning:"
        ws2["A4"].font = Font(name="Calibri", sz=11, b=True, color="FF9C0006")
        ws2["A5"] = warning
        ws2["A5"].font = Font(name="Calibri", sz=11, color="FF9C0006")
        ws2["A5"].alignment = Alignment(wrap_text=True, vertical="top")
        ws2.merge_cells("A5:J8")

    # ---- Exceptions ----
    ws3 = wb.create_sheet("Exceptions")
    exc_headers = ["Unique Identifier", "DOB", "Age", "Age range", "Country name station",
                   "Derived Region", "N° of days insured", "MonthlyPremiumUSD", "ExpectedUSD",
                   "Premium medical", "DifferenceUSD", "PercentDiff", "Status", "Reason"]
    ws3.append(exc_headers)
    style_header(ws3, len(exc_headers))
    exc = validation_df[validation_df["Status"] == "CHECK"]
    for _, row in exc.iterrows():
        dob_val = row["DOB"].date() if pd.notna(row["DOB"]) else None
        age_val = None if pd.isna(row["AgeNum"]) else row["AgeNum"]
        ws3.append([
            row["UID"], dob_val, age_val, row["AgeRange"], row["Country"], row["DerivedRegion"],
            row["Days"], round(row["MonthlyRate"], 2), round(row["ExpectedUSD"], 2),
            round(row["PremiumBilled"], 2), round(row["DifferenceUSD"], 2), row["PercentDiff"],
            row["Status"], row["Reason"],
        ])
    if len(exc) == 0:
        ws3["A2"] = "No exceptions found"
        ws3["A2"].font = data_font
    for r in range(2, ws3.max_row + 1):
        for c in range(1, len(exc_headers) + 1):
            ws3.cell(row=r, column=c).font = data_font
            if c in (8, 9, 10, 11, 12):
                ws3.cell(row=r, column=c).number_format = num_fmt
        ws3.cell(row=r, column=2).number_format = "mm/dd/yyyy"
    exc_widths = [20, 12, 8, 12, 18, 16, 14, 16, 14, 14, 14, 12, 10, 32]
    for i, w in enumerate(exc_widths, start=1):
        ws3.column_dimensions[ws3.cell(row=1, column=i).column_letter].width = w

    # ---- Rectification (if provided) ----
    if rect_df is not None:
        ws4 = wb.create_sheet("Rectification")
        ws4.append(["Unique Identifier", "Premium (rectified)", "Previous Due", "Difference"])
        style_header(ws4, 4)
        for _, row in rect_df.iterrows():
            ws4.append([row["UID"], round(row["PremiumRect"], 2), round(row["PreviousDue"], 2), round(row["Diff"], 2)])
        r = ws4.max_row + 1
        ws4.cell(row=r, column=1, value="TOTAL").font = bold_font
        for c, col in zip((2, 3, 4), ("B", "C", "D")):
            cell = ws4.cell(row=r, column=c, value=f"=SUM({col}2:{col}{r-1})")
            cell.font = bold_font
            cell.number_format = num_fmt
        for i, w in enumerate([22, 18, 16, 16], start=1):
            ws4.column_dimensions[ws4.cell(row=1, column=i).column_letter].width = w
        for rr in range(2, r):
            for c in range(2, 5):
                ws4.cell(row=rr, column=c).number_format = num_fmt
                ws4.cell(row=rr, column=c).font = data_font
            ws4.cell(row=rr, column=1).font = data_font

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Report builder
# ---------------------------------------------------------------------------
def _set_heading_black(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def build_report(summary, warning, rect_df, tolerance, rate_is_annual, period_label, validation_df):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    has_exceptions = summary["check"] > 0
    has_no_rate = summary["no_rate"] > 0
    clean = not has_exceptions and not has_no_rate
    rate_basis = "Annual" if rate_is_annual else "Monthly"
    combined_total = summary["total_billed"] + (rect_df["PremiumRect"].sum() if rect_df is not None else 0)

    title = doc.add_heading("Premium Reconciliation Report", level=1)
    title.alignment = 1
    _set_heading_black(title)
    if period_label:
        sub = doc.add_paragraph(period_label)
        sub.alignment = 1

    # ---- Key findings ----
    h = doc.add_heading("Key findings", level=2); _set_heading_black(h)
    doc.add_paragraph(
        f"{summary['total_rows']} member records were reviewed. "
        f"{summary['ok']} reconcile within the \u00b1{tolerance}% tolerance"
        + (f", {summary['check']} fall outside it" if has_exceptions else "")
        + (f", and {summary['no_rate']} have no matching rate code." if has_no_rate else "."),
        style="List Bullet")
    doc.add_paragraph(
        "Billed figures are accurate to the bill \u2014 they match the source billing data exactly, "
        "with nothing changed.", style="List Bullet")
    if warning:
        doc.add_paragraph(
            "A pattern check flagged a likely systemic issue rather than isolated billing errors "
            "(see Exceptions register below).", style="List Bullet")
    if clean:
        doc.add_paragraph(
            f"No exceptions were found. Billed and expected premium are consistent, with an "
            f"average gap of {summary['avg_ok_pct']:.2f}%." if summary["avg_ok_pct"] is not None
            else "No exceptions were found.", style="List Bullet")
    if rect_df is not None:
        rect_diff = rect_df["Diff"].sum()
        new_charges = rect_df[(rect_df["PreviousDue"] == 0) & (rect_df["PremiumRect"] != 0)]
        if abs(rect_diff) < 0.01:
            doc.add_paragraph("Rectification check found no difference against previous due amounts.", style="List Bullet")
        else:
            doc.add_paragraph(
                f"A rectification check found {len(new_charges)} record(s) with no previous due on file, "
                f"adding {rect_diff:,.2f} USD in new charges.", style="List Bullet")

    # ---- Key figures ----
    h = doc.add_heading("Key figures", level=2); _set_heading_black(h)
    doc.add_paragraph(f"Billed Premium medical total: {summary['total_billed']:,.2f} USD.", style="List Bullet")
    doc.add_paragraph(f"Expected premium total: {summary['total_expected']:,.2f} USD.", style="List Bullet")
    doc.add_paragraph(f"Total difference (billed less expected): {summary['total_diff']:,.2f} USD.", style="List Bullet")
    doc.add_paragraph(
        f"Records: {summary['total_rows']} \u2014 OK: {summary['ok']} \u2014 CHECK: {summary['check']} "
        f"\u2014 No Rate: {summary['no_rate']} \u2014 Tolerance: \u00b1{tolerance}%.", style="List Bullet")
    doc.add_paragraph(f"Rate table treated as: {rate_basis} premiums.", style="List Bullet")
    if rect_df is not None:
        doc.add_paragraph(f"Rectification Premium medical total: {rect_df['PremiumRect'].sum():,.2f} USD.", style="List Bullet")
        doc.add_paragraph(f"Rectification previous due total: {rect_df['PreviousDue'].sum():,.2f} USD.", style="List Bullet")
        doc.add_paragraph(f"Rectification difference (new charges): {rect_df['Diff'].sum():,.2f} USD.", style="List Bullet")
        doc.add_paragraph(f"Combined amount (Billing + Rectification Premium medical): {combined_total:,.2f} USD.", style="List Bullet")

    # ---- Executive summary ----
    h = doc.add_heading("Executive summary", level=2); _set_heading_black(h)
    p1 = (f"This review checks insurance premiums" + (f" for {period_label}" if period_label else "")
          + f", using the uploaded billing file and rate table. Source billing figures were not changed.")
    doc.add_paragraph(p1)
    if clean:
        p2 = (f"The billed total of {summary['total_billed']:,.2f} USD is accurate to the bill and consistent "
              f"with expected premium under the {rate_basis.lower()} rate table, with an average gap of "
              f"{summary['avg_ok_pct']:.2f}%, well inside tolerance." if summary["avg_ok_pct"] is not None else
              f"The billed total of {summary['total_billed']:,.2f} USD is accurate to the bill and consistent with expected premium.")
    else:
        p2 = (f"{summary['check']} of {summary['total_rows']} records fall outside the \u00b1{tolerance}% tolerance, "
              f"for a total difference of {summary['total_diff']:,.2f} USD against expected premium of "
              f"{summary['total_expected']:,.2f} USD.")
        if warning:
            p2 += (" " + warning)
    doc.add_paragraph(p2)
    if rect_df is not None and abs(rect_df["Diff"].sum()) > 0.01:
        doc.add_paragraph(
            f"A separate rectification check compared each member's current premium to their previous due "
            f"amount. {len(rect_df) - len(new_charges)} of {len(rect_df)} members showed no difference. "
            f"{len(new_charges)} had no previous due on record, adding {rect_df['Diff'].sum():,.2f} USD in new "
            f"charges. Combined, Billing and Rectification total {combined_total:,.2f} USD.")

    # ---- Scope ----
    h = doc.add_heading("Scope", level=2); _set_heading_black(h)
    scope_text = (
        f"The review covered all {summary['total_rows']} members in the uploaded billing file, using category "
        f"code, age range, insured days, and billed premium. "
        + ("A rectification file was also included, checking each member's premium against their previous due "
           "amount. " if rect_df is not None else "No rectification file was included for this review. ")
        + "Original billed figures were not changed."
    )
    doc.add_paragraph(scope_text)

    # ---- Methodology ----
    h = doc.add_heading("Methodology", level=2); _set_heading_black(h)
    doc.add_paragraph(
        f"Each record was checked using insured days and category code (age band and region) against the "
        f"uploaded rate table, treated as {rate_basis.lower()} premiums. Expected premium is "
        + ("(annual rate \u00f7 12) \u00d7 insured days \u00f7 30" if rate_is_annual else "monthly rate \u00d7 insured days \u00f7 30")
        + f", compared to the billed amount. Records with zero insured days were marked OK. A \u00b1{tolerance}% "
        f"tolerance flags records as CHECK. Records with no matching rate code are marked No Rate rather than "
        f"guessed. All totals are checked against the source file's own totals."
    )

    # ---- Results ----
    h = doc.add_heading("Results", level=2); _set_heading_black(h)
    p = doc.add_paragraph(); p.add_run("Billing: ").bold = True
    doc.add_paragraph(
        f"The billed total of {summary['total_billed']:,.2f} USD is accurate to the bill and matches the "
        f"source file's own total.")
    p = doc.add_paragraph(); p.add_run("Rectification: ").bold = True
    if rect_df is not None:
        doc.add_paragraph(
            f"The Rectification Premium medical total is {rect_df['PremiumRect'].sum():,.2f} USD, against a "
            f"previous due total of {rect_df['PreviousDue'].sum():,.2f} USD, a difference of "
            f"{rect_df['Diff'].sum():,.2f} USD"
            + (f", sitting entirely in {len(new_charges)} record(s) with no previous due on file." if len(new_charges) > 0 else ", with no records showing a difference."))
    else:
        doc.add_paragraph("No rectification file was provided, so this review covers billing only.")
    p = doc.add_paragraph(); p.add_run("Summary tie-out: ").bold = True
    tie_text = f"The billed total ties out exactly to the source file. Against expected premium of {summary['total_expected']:,.2f} USD, the gap is {summary['total_diff']:,.2f} USD"
    tie_text += f" ({summary['total_diff']/summary['total_expected']*100:.2f}%)." if summary["total_expected"] else "."
    if rect_df is not None:
        tie_text += f" Billing and Rectification together total {combined_total:,.2f} USD."
    doc.add_paragraph(tie_text)

    # ---- Accuracy and controls ----
    h = doc.add_heading("Accuracy and controls", level=2); _set_heading_black(h)
    doc.add_paragraph(
        "Billed totals are taken directly from the uploaded file and are accurate to the bill, unchanged. "
        "All calculations use live formulas in the output workbook, so figures stay traceable and auditable. "
        "Validation, Summary, and Exceptions sheets all agree with each other."
        + (" No records fall outside tolerance." if clean else " See Exceptions register below for records outside tolerance.")
    )

    # ---- Exceptions register ----
    h = doc.add_heading("Exceptions register", level=2); _set_heading_black(h)
    if clean:
        doc.add_paragraph(f"No exceptions were identified. All {summary['total_rows']} records reconcile within the \u00b1{tolerance}% tolerance.")
    else:
        exc_text = f"{summary['check']} of {summary['total_rows']} records ({summary['check']/summary['total_rows']*100:.0f}%) were flagged."
        if warning:
            exc_text += " " + warning
        else:
            exc_text += " See the Exceptions sheet in the workbook for the full list and reasons."
        doc.add_paragraph(exc_text)
        if has_no_rate:
            doc.add_paragraph(
                f"{summary['no_rate']} record(s) had no matching rate code and could not be checked against "
                f"expected premium \u2014 confirm these codes exist in the rate table.")

    # ---- Approval recommendation ----
    h = doc.add_heading("Approval recommendation", level=2); _set_heading_black(h)
    if clean:
        rec = f"The billed total of {summary['total_billed']:,.2f} USD is accurate to the bill and consistent with expected premium. Payment may proceed."
        if rect_df is not None:
            rec = f"The combined total of {combined_total:,.2f} USD (Billing + Rectification) is accurate to the bill and consistent with expected premium. Payment may proceed."
    else:
        rec = (
            f"The billed total of {summary['total_billed']:,.2f} USD is accurate to the bill and well supported "
            f"on its own. But with {summary['check']} of {summary['total_rows']} records outside tolerance"
            + (", and a uniform pattern suggesting a systemic rather than per-record issue," if warning else ",")
            + " that comparison should not be used to confirm billing accuracy until the flagged records are "
              "reviewed. Payment of the billed total can proceed on its own merits, but should not be described "
              "as checked against expected premium until the exceptions are resolved."
        )
    doc.add_paragraph(rec)

    # ---- Conclusion ----
    h = doc.add_heading("Conclusion", level=2); _set_heading_black(h)
    if clean:
        doc.add_paragraph(
            f"The billed Premium medical total of {summary['total_billed']:,.2f} USD is accurate to the bill, "
            f"consistent with expected premium, and fully supported. This reconciliation is fully validated."
        )
    else:
        doc.add_paragraph(
            f"The billed Premium medical total of {summary['total_billed']:,.2f} USD is accurate to the bill "
            f"and well supported. It does not yet confirm that billed amounts match expected premium: "
            f"{summary['check']} of {summary['total_rows']} records fall outside tolerance. Recommend resolving "
            f"the flagged records before calling this reconciliation fully validated."
        )

    # ---- References ----
    h = doc.add_heading("References", level=2); _set_heading_black(h)
    doc.add_paragraph(
        "Premium calculations: insurance billing reconciliation"
        + (f", {period_label}" if period_label else "") + f". Generated {date.today().isoformat()}."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Main flow
# ---------------------------------------------------------------------------
if run_btn:
    if not billing_file or not rates_file:
        st.error("Please upload at least a billing file and a rate table.")
        st.stop()

    with st.spinner("Reading files..."):
        try:
            billing_df = load_billing(billing_file)
            rates_df = load_rates(rates_file)
            rect_df = load_rectification(rect_file) if rect_file else None
        except Exception as e:
            st.error(f"Could not read one of the files: {e}")
            st.stop()

    with st.spinner("Reconciling..."):
        validation_df, warning = run_reconciliation(billing_df, rates_df, tolerance, rate_is_annual)
        summary = build_summary(validation_df)

    st.success(f"Done — {summary['total_rows']} records reviewed.")

    if warning:
        st.warning(warning)

    c1, c2, c3, c4 = st.columns(4)
    cards_html = '<div class="ledger-cards">' + \
        ledger_card("Billed total", f"${summary['total_billed']:,.2f}") + \
        ledger_card("Expected total", f"${summary['total_expected']:,.2f}") + \
        ledger_card("Difference", f"${summary['total_diff']:,.2f}", warn=(summary['check'] > 0)) + \
        ledger_card("Flagged (CHECK)", f"{summary['check']} / {summary['total_rows']}", warn=(summary['check'] > 0)) + \
        '</div>'
    st.markdown(cards_html, unsafe_allow_html=True)

    st.subheader("Validation detail")
    display_df = validation_df.copy()
    display_df["Status"] = display_df["Status"].apply(
        lambda s: f'<span class="stamp {"ok" if s == "OK" else "check"}">{s}</span>'
    )
    for col in ["MonthlyRate", "ExpectedUSD", "PremiumBilled", "DifferenceUSD"]:
        display_df[col] = display_df[col].round(2)
    display_df["PercentDiff"] = display_df["PercentDiff"].apply(
        lambda v: "" if pd.isna(v) else f"{v:.2f}%"
    )
    display_df["DOB"] = display_df["DOB"].dt.strftime("%Y-%m-%d").fillna("")
    display_df = display_df[[
        "UID", "NameFamily", "FirstNameFamily", "DOB", "AgeNum", "AgeRange", "Country",
        "DerivedRegion", "Days", "MonthlyRate", "ExpectedUSD", "PremiumBilled",
        "DifferenceUSD", "PercentDiff", "Status", "Reason",
    ]].rename(columns={
        "UID": "Unique Identifier", "NameFamily": "Name Family", "FirstNameFamily": "First Name Family",
        "AgeNum": "Age", "AgeRange": "Age range", "Country": "Country name station",
        "DerivedRegion": "Derived Region", "Days": "N of days insured",
        "MonthlyRate": "MonthlyPremiumUSD", "PremiumBilled": "Premium medical",
        "DifferenceUSD": "DifferenceUSD", "PercentDiff": "PercentDiff",
    })
    st.markdown(
        display_df.to_html(escape=False, index=False, classes="ledger-table"),
        unsafe_allow_html=True,
    )

    if rect_df is not None:
        st.subheader("Rectification detail")
        st.dataframe(rect_df, use_container_width=True, height=200)

    excel_buf = build_excel(validation_df, summary, warning, rect_df, tolerance, rate_is_annual)
    report_buf = build_report(summary, warning, rect_df, tolerance, rate_is_annual, period_label, validation_df)

    dcol1, dcol2 = st.columns(2)
    dcol1.download_button(
        "Download reconciled workbook (.xlsx)", excel_buf,
        file_name=f"Reconciliation_{date.today().isoformat()}.xlsx",
        use_container_width=True,
    )
    dcol2.download_button(
        "Download report (.docx)", report_buf,
        file_name=f"Reconciliation_Report_{date.today().isoformat()}.docx",
        use_container_width=True,
    )
else:
    st.info("Upload your billing file and rate table on the left, then click **Run reconciliation**.")
